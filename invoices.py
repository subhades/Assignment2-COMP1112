"""
Amanda Robinson
1181969
Assignment 2
COMP1112

This program is for the COMP 1112 assignment 2 due March 23rd 2023 
It will grab all the information from 200 different word documents and store the data it contains appropriately in 
an excel sheet that will looks simiarly to the A2_Ex.xlsx example provided to us."""

import re
import os
import docx
#import linecache
import openpyxl
from openpyxl import Workbook

#defining my path to the invoices 
doc_dir = 'invoices'
#The numeric pattern I will be using to locate the numbers in a line of text
numeric_pattern = re.compile(r'\d+\.\d+|\d+')
#create a new excel workbook
invoice_list = openpyxl.Workbook()
#create a new excel worksheet
invoice_data = invoice_list.active
#title my worksheet 
invoice_data.title = "invoices - Assignment 2"
#Create my initial title cells related to the information they will hold
invoice_data.cell(row=1, column=1, value="Invouice Number")
invoice_data.cell(row=1, column=2, value="Number of products")
invoice_data.cell(row=1, column=3, value="Subtotal")
invoice_data.cell(row=1, column=4, value="Tax")
invoice_data.cell(row=1, column=5, value="Total")

#Keep track of my loops with int to help with assigning cells.
i=0

#for loop will loop through each document and read the name in the directory I made earlier
for filename in os.listdir(doc_dir):
    #if statement that states if the document is a docx word document (ends in .docx) to enter the loop
    if filename.endswith('.docx'):
        #To keep track of the loop count for the cells. add one each time we enter loop with a new document 
        i=i+1
        
        #open the path and name the path
        doc_path = os.path.join(doc_dir, filename)
        #assign the document to the variable doc
        doc = docx.Document(doc_path)
        # line_count = 0 no longer relevant 

        #I figured out that the document was divided into three paragraphs
        #pull invoice numbers from the 1st paragraph 
        target_invoices = doc.paragraphs[0]
        #remove the entire paragraph as it only contains the one line and store it in the variable invoice
        invoice = target_invoices.text.strip()
        #I print it a lot as I test to make sure it is grabbing the information I want 
        print(invoice)
        #add the invoice as the value in the correct row (one + the cycle number it is in because we don't want overlapping + we need the first row to be the title row )
        invoice_data.cell(row=1+i, column=1, value=invoice)

        #turning the second paragraph containing all the product information nto a variable
        #I want to return only the total number of products and not the products themselves
        target_product_count = doc.paragraphs[1]
        #removing white spaces with the strip() function and pulling the text
        products = target_product_count.text.strip()
        #using the re.findeall function to find all numbers in the paragraph that contain numbers
        numbers_products = re.findall(r'\d+', products)
        #finding the total of all the products in the paragraph 
        total_product = sum([int(num) for num in numbers_products])
        #printing again for my own benefit 
        print("Product amount: ", total_product)
        #storing in a similar fashion invoices, but in the corresponding column
        invoice_data.cell(row=1+i, column=2, value=total_product)

        #store the third paragraph into a variable 
        target_price_count = doc.paragraphs[2]
        #storing the three values and removing white space into an list 
        subtotal_string = target_price_count.runs[0].text.split()
        #print to make sure it worked
        print(subtotal_string)

        #x3 I used re.findall (although it was simply one number) because my original method did work, but the excel sheet was 
        #very angry and was telling me that it was being stored as text instead of a number. But I found all(the single number) based
        #on the pattern created at the top and went through each item to pull the appropriate number
        #and made sure it was a float variable. (re.search and .group())
        subtotal = re.findall(numeric_pattern, subtotal_string[0])
        subtotal = sum([float(num) for num in subtotal])

        tax = re.findall(numeric_pattern, subtotal_string[1])
        tax = sum([float(num) for num in tax])

        total = re.findall(numeric_pattern, subtotal_string[2])
        total = sum([float(num) for num in total])

        #three if statements that print the values if the program found a value that matches numeric pattern established and then 
        # adds it to the excel sheet like above three 
        if subtotal:
            print(subtotal)
            invoice_data.cell(row=1+i, column=3, value=subtotal)
        if tax:
            print(tax)
            invoice_data.cell(row=1+i, column=4, value=tax)
        if total:
            print(total)
            invoice_data.cell(row=1+i, column=5, value=total)

    #Save the excel workbook
    invoice_list.save(filename="A2_1181969.xlsx")
